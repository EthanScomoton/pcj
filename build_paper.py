"""
LNEE Springer-Format Paper Generator
=====================================
Builds a .docx file for the Port Integrated Energy System paper, strictly
following the formatting rules from
"Paper Format/LNEE Springer Guidelines forAuthors of Proceedings".

Currently included sections (per user request):
  1. Title, authors, affiliation, abstract, keywords
  2. Section 1 — Introduction (overview / literature review / contributions)
  3. Section 2 — System Description and Modelling
     (Includes the pre-trained .pth deep-learning forecasting model)

NOT yet included: Section 3 Results, Section 4 Conclusion.

Mathematical notation is rendered with Word's native subscript / superscript
runs so that variable indices look correct in Word (without relying on the
incomplete Unicode subscript block which only covers a few Latin letters).
"""
from __future__ import annotations

import os

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# =============================================================================
# Style helpers
# =============================================================================
FONT_LATIN = 'Times New Roman'


def _set_run_font(run, name=FONT_LATIN, size_pt=10, bold=False, italic=False,
                  sub=False, sup=False):
    """Apply Times New Roman to a run (and override East-Asian font tag)."""
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.subscript = sub
    run.font.superscript = sup
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:cs'), name)


# -----------------------------------------------------------------------------
# Mixed-formatting paragraph (supports subscripts/superscripts inline)
# -----------------------------------------------------------------------------
def add_mixed(doc, segments, *, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              first_line_indent_cm=0.5, space_before_pt=0, space_after_pt=2,
              line_spacing=1.18, bold_default=False, italic_default=False):
    """Add a paragraph composed of (text, format_options) segments.

    Each segment is either a string or a tuple/dict with keys:
      'text'   — string
      'bold'   — bool (default False)
      'italic' — bool (default False)
      'sub'    — bool
      'sup'    — bool
      'size'   — pt (default = paragraph size)
    """
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    pf.line_spacing = line_spacing
    if first_line_indent_cm > 0 and align != WD_ALIGN_PARAGRAPH.CENTER:
        pf.first_line_indent = Cm(first_line_indent_cm)

    for seg in segments:
        if isinstance(seg, str):
            text = seg
            opts = {}
        else:
            text = seg['text']
            opts = seg
        run = p.add_run(text)
        _set_run_font(
            run,
            size_pt=opts.get('size', size),
            bold=opts.get('bold', bold_default),
            italic=opts.get('italic', italic_default),
            sub=opts.get('sub', False),
            sup=opts.get('sup', False),
        )
    return p


def add_para(doc, text='', *, size=10, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             first_line_indent_cm=0.0,
             space_before_pt=0, space_after_pt=2,
             line_spacing=1.15):
    """Simple-text paragraph helper."""
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    pf.line_spacing = line_spacing
    if first_line_indent_cm > 0:
        pf.first_line_indent = Cm(first_line_indent_cm)
    if text:
        run = p.add_run(text)
        _set_run_font(run, size_pt=size, bold=bold, italic=italic)
    return p


def add_title(doc, text):
    return add_para(doc, text, size=14, bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before_pt=12, space_after_pt=6,
                    line_spacing=1.2)


def add_h1(doc, number, text):
    return add_para(doc, f"{number}\u2003{text}", size=12, bold=True,
                    align=WD_ALIGN_PARAGRAPH.LEFT,
                    space_before_pt=12, space_after_pt=4, line_spacing=1.15)


def add_h2(doc, number, text):
    return add_para(doc, f"{number}\u2003{text}", size=10, bold=True,
                    align=WD_ALIGN_PARAGRAPH.LEFT,
                    space_before_pt=8, space_after_pt=2, line_spacing=1.15)


def add_body(doc, text, *, indent=True):
    return add_para(doc, text, size=10,
                    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    first_line_indent_cm=0.5 if indent else 0.0,
                    space_after_pt=2, line_spacing=1.18)


# -----------------------------------------------------------------------------
# Equation builders
# -----------------------------------------------------------------------------
def add_equation(doc, segments, number, *, size=10):
    """
    Add a centered equation with a right-aligned equation number.
    `segments` is a list of (text, formatting-options) like `add_mixed`,
    so subscripts / superscripts / italics render properly in Word.
    """
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Cm(13.0)
    tbl.columns[1].width = Cm(2.0)

    cell_eq = tbl.cell(0, 0)
    p_eq = cell_eq.paragraphs[0]
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_eq.paragraph_format.space_before = Pt(2)
    p_eq.paragraph_format.space_after = Pt(2)
    for seg in segments:
        if isinstance(seg, str):
            text, opts = seg, {}
        else:
            text, opts = seg['text'], seg
        run = p_eq.add_run(text)
        _set_run_font(
            run,
            size_pt=opts.get('size', size),
            bold=opts.get('bold', False),
            italic=opts.get('italic', False),
            sub=opts.get('sub', False),
            sup=opts.get('sup', False),
        )

    cell_no = tbl.cell(0, 1)
    p_no = cell_no.paragraphs[0]
    p_no.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_no.paragraph_format.space_before = Pt(2)
    p_no.paragraph_format.space_after = Pt(2)
    r_no = p_no.add_run(f"({number})")
    _set_run_font(r_no, size_pt=size)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return tbl


def add_fig_caption(doc, label, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(6)
    r_bold = p.add_run(f"{label}. ")
    _set_run_font(r_bold, size_pt=9, bold=False)
    r_text = p.add_run(text)
    _set_run_font(r_text, size_pt=9)
    return p


def add_table_caption(doc, label, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(2)
    r_label = p.add_run(f"{label}. ")
    _set_run_font(r_label, size_pt=9)
    r_text = p.add_run(text)
    _set_run_font(r_text, size_pt=9)
    return p


def add_abstract(doc, abstract_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.15
    r_label = p.add_run("Abstract.\u2002")
    _set_run_font(r_label, size_pt=9, bold=True)
    r_body = p.add_run(abstract_text)
    _set_run_font(r_body, size_pt=9)
    return p


def add_keywords(doc, keywords):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(10)
    pf.line_spacing = 1.15
    r_label = p.add_run("Keywords:\u2002")
    _set_run_font(r_label, size_pt=9, bold=True)
    r_text = p.add_run(" \u00b7 ".join(keywords))
    _set_run_font(r_text, size_pt=9)
    return p


def add_authors(doc, authors_line):
    return add_para(doc, authors_line, size=10,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before_pt=4, space_after_pt=2, line_spacing=1.15)


def add_affiliation(doc, line):
    return add_para(doc, line, size=9, italic=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before_pt=0, space_after_pt=2, line_spacing=1.15)


def add_email(doc, email):
    return add_para(doc, email, size=9,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before_pt=0, space_after_pt=8, line_spacing=1.15)


# =============================================================================
# Configure document layout (Springer LNCS/LNEE 152 × 235 mm trim size)
# =============================================================================
def configure_document(doc):
    section = doc.sections[0]
    section.page_height = Cm(23.5)
    section.page_width = Cm(15.2)
    section.top_margin = Cm(2.6)
    section.bottom_margin = Cm(2.6)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    style = doc.styles['Normal']
    style.font.name = FONT_LATIN
    style.font.size = Pt(10)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), FONT_LATIN)
    rFonts.set(qn('w:hAnsi'), FONT_LATIN)
    rFonts.set(qn('w:eastAsia'), FONT_LATIN)
    rFonts.set(qn('w:cs'), FONT_LATIN)


# =============================================================================
# Inline mixed-format helpers (used inside body paragraphs that contain math)
# =============================================================================
def s(text, **kwargs):
    """Shorthand for a formatted segment."""
    return {'text': text, **kwargs}


# Common short-hand formatters for inline math snippets
def italic(text):  return {'text': text, 'italic': True}
def sub(text):     return {'text': text, 'sub': True}
def sup(text):     return {'text': text, 'sup': True}


# =============================================================================
# Build the paper
# =============================================================================
def build_paper(output_path):
    doc = Document()
    configure_document(doc)

    # ---------------------------------------------------------------------
    # Title block
    # ---------------------------------------------------------------------
    add_title(
        doc,
        "Deep-Learning Load Forecasting and Conformal-Robust "
        "Model Predictive Control for Carbon-Aware Energy Management "
        "of Dry-Bulk Port Integrated Energy Systems"
    )

    add_authors(doc, "First Author\u00b9(\u2709), Second Author\u00b9, and Third Author\u00b2")
    add_affiliation(doc,
        "\u00b9 College of Electrical and Information Engineering, "
        "Hunan University, Changsha 410082, China")
    add_affiliation(doc,
        "\u00b2 Department of Electrical Engineering, "
        "Tsinghua University, Beijing 100084, China")
    add_email(doc, "corresponding.author@hnu.edu.cn")

    # ---------------------------------------------------------------------
    # Abstract & Keywords
    # ---------------------------------------------------------------------
    add_abstract(doc,
        "Dry-bulk ports are highly energy-intensive nodes of the global "
        "logistic network, and their decarbonisation under the dual-carbon "
        "policy requires the joint optimisation of distributed photovoltaic "
        "and wind generation, battery energy storage, and grid procurement. "
        "However, the stochastic nature of port loads and renewable output "
        "makes single-point forecast-driven dispatch fragile, while purely "
        "deterministic model predictive control cannot expose the residual "
        "forecast risk to the decision layer. This paper proposes an "
        "integrated energy-management framework that couples a pre-trained "
        "deep neural load-forecasting model with calibrated distribution-free "
        "uncertainty quantification and a parametrised carbon-aware robust "
        "model predictive controller. The forecasting backbone is a "
        "bidirectional long short-term memory network augmented with a "
        "feature-gating mechanism and dual (temporal\u2013feature) attention, "
        "whose pre-trained weights (best_EModel_FeatureWeight4.pth) are "
        "loaded at run time for inference; a split-conformal calibrator "
        "subsequently converts the point forecasts into prediction intervals "
        "with finite-sample marginal coverage guarantee. The same forecast "
        "trajectory is shared across five dispatch strategies\u2014a "
        "no-storage baseline, a rule-based peak-shaving heuristic, an "
        "economic MPC, a carbon-aware MPC, and a conformal robust "
        "MPC\u2014all implemented through a parametrised CVXPY problem "
        "that is built once and warm-started at each rolling step. "
        "Regional grid emission factors and an hour-by-hour dynamic "
        "marginal carbon-intensity profile are embedded in the "
        "multi-objective cost function to enable explicit "
        "cost\u2013carbon trade-off. The present paper documents the system "
        "architecture, mathematical formulation and forecasting model that "
        "constitute the proposed framework; quantitative evaluation results "
        "and conclusions will be reported in subsequent work.")

    add_keywords(doc, [
        "Port integrated energy system",
        "Deep-learning load forecasting",
        "Conformal prediction",
        "Model predictive control",
        "Carbon-aware dispatch",
        "Battery energy storage",
    ])

    # =====================================================================
    # 1. Introduction
    # =====================================================================
    add_h1(doc, "1", "Introduction")

    add_body(doc,
        "The decarbonisation of port operations has become a central "
        "objective of the global maritime industry. Coastal ports account "
        "for an increasingly significant share of national industrial "
        "electricity consumption, and dry-bulk terminals\u2014where "
        "high-power ship loaders, conveyor belts, refrigeration warehouses "
        "and onshore-power facilities operate in parallel\u2014are among "
        "the most carbon-intensive logistic nodes. Under the dual-carbon "
        "policy framework recently endorsed at national level, the "
        "transition from grid-dominated supply to a port integrated energy "
        "system (PIES) that couples photovoltaic (PV) and wind generation, "
        "battery energy storage (BES) and the public grid has emerged as "
        "the principal technical pathway to reduce both the operating cost "
        "and the CO\u2082 footprint of port enterprises [1\u20133].",
        indent=False)

    add_body(doc,
        "Although a number of recent studies have demonstrated the "
        "techno-economic potential of PIES, two coupled difficulties still "
        "stand in the way of practical deployment. First, port loads "
        "exhibit strong non-stationarity and event-driven peaks: berth "
        "occupancy, ship-to-shore (STS) crane scheduling and refrigeration "
        "duty cycles introduce dynamics that classical statistical "
        "predictors handle poorly. Second, the high penetration of "
        "intermittent renewable generation propagates additional "
        "uncertainty into the dispatch problem, so that any deterministic "
        "scheduling decision based solely on a single-point load forecast "
        "can become economically inefficient or even infeasible when the "
        "realisation deviates from the predicted trajectory.")

    add_body(doc,
        "Three streams of literature address these difficulties in "
        "isolation. (i) Data-driven load forecasting has matured rapidly: "
        "recurrent neural networks, in particular long short-term memory "
        "(LSTM) variants, have been shown to outperform autoregressive "
        "models on industrial demand series [4, 5]; subsequent works "
        "introduce attention mechanisms to focus on the most informative "
        "time steps [6] and feature-gating modules to reweight noisy "
        "exogenous inputs such as ambient temperature, humidity and "
        "renewable generation [7]. (ii) Uncertainty quantification for "
        "energy forecasts has evolved from Bayesian deep ensembles and "
        "Monte-Carlo dropout to conformal prediction, a distribution-free "
        "framework that provides finite-sample marginal coverage "
        "guarantees and that, importantly, can be applied as a post-hoc "
        "wrapper around any pre-trained model without retraining [8, 9]. "
        "(iii) Energy management of grid-connected storage has moved from "
        "rule-based peak-shaving to model predictive control (MPC) with "
        "rolling-horizon optimisation [10, 11], and more recently to "
        "carbon-aware MPC formulations in which marginal grid emission "
        "factors enter the objective function alongside the electricity "
        "tariff [12, 13].")

    add_body(doc,
        "Despite this progress, a unified framework that simultaneously "
        "(a) exploits an industrial-grade deep neural load predictor, "
        "(b) propagates calibrated uncertainty from the predictor into "
        "the controller, and (c) optimises an explicit "
        "cost\u2013carbon\u2013peak trade-off on real port data is still "
        "missing. Most existing PIES studies either feed a single-point "
        "neural forecast directly into MPC, thereby ignoring residual "
        "predictor error, or invoke chance constraints under restrictive "
        "Gaussian-noise assumptions that are routinely violated by the "
        "non-stationary load of dry-bulk ports. Furthermore, the "
        "computational burden of repeatedly building large convex "
        "programmes over an 8 760-hour annual horizon is rarely addressed, "
        "limiting the applicability of MPC-based studies to short case "
        "studies of a few representative days.")

    add_body(doc,
        "Motivated by these gaps, the present work proposes and describes "
        "a deep-learning-driven, conformal-robust, carbon-aware model "
        "predictive energy-management framework for dry-bulk port "
        "integrated energy systems. The main contributions are summarised "
        "as follows:")

    add_body(doc,
        "(1) A bidirectional LSTM forecasting model with feature gating "
        "and dual (temporal\u2013feature) attention is integrated into the "
        "dispatch pipeline; its pre-trained weights "
        "(best_EModel_FeatureWeight4.pth) are loaded at run time, so "
        "that no online re-training is required and the framework can be "
        "deployed on standard industrial hardware.",
        indent=False)

    add_body(doc,
        "(2) A split-conformal calibration layer is wrapped around the "
        "pre-trained predictor to convert each point forecast into a "
        "prediction interval with a target marginal coverage of "
        "1\u2009\u2212\u2009\u03b1; the same calibrated interval is "
        "subsequently used as the worst-case demand bound for the robust "
        "controller, ensuring that the operational risk is bounded by an "
        "explicit, user-tunable confidence level.",
        indent=False)

    add_body(doc,
        "(3) A parametrised rolling-horizon MPC core is constructed once "
        "and re-solved with warm starts and a three-tier solver chain "
        "(SCS \u2192 CLARABEL \u2192 OSQP) with a heuristic fallback, "
        "lowering the per-step solve time so that year-long (8 760-hour) "
        "studies are computationally tractable; five strategies share the "
        "same controller skeleton and the same forecast cache, "
        "eliminating redundant forward passes of the neural network.",
        indent=False)

    add_body(doc,
        "(4) Regional grid emission factors for the six Chinese "
        "interconnection regions and an hour-resolved dynamic marginal "
        "carbon-intensity profile are embedded as explicit decision "
        "variables in the objective function, enabling a quantitative "
        "trade-off between electricity cost, peak demand and CO\u2082 "
        "emission, and supporting Pareto-front analysis across competing "
        "operating regimes.",
        indent=False)

    add_body(doc,
        "The remainder of this paper is organised as follows. Section 2 "
        "presents the architecture of the proposed port integrated "
        "energy-management framework, formulates the energy balance and "
        "the battery storage dynamics, details the deep neural "
        "load-forecasting backbone together with the conformal calibration "
        "layer, and develops the parametrised carbon-aware rolling-horizon "
        "controller. Quantitative results, ablation studies and concluding "
        "remarks will be reported subsequently.")

    # =====================================================================
    # 2. System Description and Modelling
    # =====================================================================
    add_h1(doc, "2", "System Description and Modelling")

    # -----------------------------------------------------------------
    # 2.1 Architecture
    # -----------------------------------------------------------------
    add_h2(doc, "2.1", "Architecture of the Port Integrated Energy System")

    add_body(doc,
        "The proposed framework follows a three-tier architecture, as "
        "illustrated in Fig.\u00a01. The data tier ingests historical load, "
        "renewable generation, ambient and calendar features at the "
        "original sub-hourly resolution and resamples them to a uniform "
        "one-hour grid by mean aggregation of numerical fields and "
        "forward-fill of categorical fields. The prediction tier hosts the "
        "pre-trained bidirectional LSTM model with feature gating and dual "
        "attention (Section\u00a02.3), together with the split-conformal "
        "calibrator (Section\u00a02.4); it consumes the most recent "
        "20-step feature window and emits a 24-hour ahead point forecast "
        "together with a calibrated prediction interval. The dispatch "
        "tier hosts five interchangeable scheduling strategies that share "
        "a common parametrised CVXPY core (Section\u00a02.5) and issue, at "
        "each hour, the battery charge/discharge and grid import "
        "set-points to be applied to the physical port equipment.",
        indent=False)

    add_fig_caption(doc, "Fig. 1",
        "Overall architecture of the proposed deep-learning-driven, "
        "conformal-robust, carbon-aware energy-management framework for "
        "dry-bulk port integrated energy systems.")

    # Body paragraph that references E_p, E_PV, etc. with proper subscripts
    add_mixed(doc, [
        "Six energy components are coupled at the port bus: the public "
        "grid feeder (with bidirectional power flow disabled by default "
        "to reflect the contractual constraints of typical industrial "
        "users), a fleet of rooftop and shed-roof photovoltaic arrays, a "
        "set of medium-voltage wind turbines installed on the pier, a "
        "lithium-iron-phosphate battery energy storage (BES) station, "
        "the aggregated port load (cranes, conveyors, refrigeration "
        "warehouses, lighting and auxiliary services), and a legacy "
        "thermal-storage discharge channel that is retained for backward "
        "compatibility with the historical operational logs. Throughout "
        "this paper, the corresponding hourly energy flows are denoted ",
        italic("E"), sub("g"), ", ",
        italic("E"), sub("PV"), ", ",
        italic("E"), sub("w"), ", ",
        italic("E"), sub("d"), ", ",
        italic("E"), sub("c"), ", ",
        italic("E"), sub("L"),
        " for the grid import, the PV utilisation, the wind utilisation, "
        "the battery discharge, the battery charge and the aggregated "
        "load, respectively.",
    ], first_line_indent_cm=0.0)

    # -----------------------------------------------------------------
    # 2.2 Energy balance and BES dynamics
    # -----------------------------------------------------------------
    add_h2(doc, "2.2", "Energy Balance and Battery Storage Dynamics")

    add_mixed(doc, [
        "At each hourly time step ",
        italic("t"), " \u2208 {1, \u2026, ", italic("T"), "}, ",
        "the aggregate port load must be served by the sum of grid "
        "purchase, renewable consumption and net storage discharge, while "
        "bidirectional flow to the upstream grid is disallowed:",
    ], first_line_indent_cm=0.0)

    add_equation(doc, [
        italic("E"), sub("L,t"), " = ",
        italic("E"), sub("g,t"), " + ",
        italic("E"), sub("PV,t"), " + ",
        italic("E"), sub("w,t"), " + ",
        italic("E"), sub("d,t"), " \u2212 ",
        italic("E"), sub("c,t"), ",    ",
        italic("E"), sub("g,t"), " \u2265 0,",
    ], "1")

    add_mixed(doc, [
        "where the renewable consumption terms are upper-bounded by the "
        "corresponding hourly forecast, that is 0 \u2264 ",
        italic("E"), sub("PV,t"), " \u2264 ",
        italic("E\u0302"), sub("PV,t"), " and 0 \u2264 ",
        italic("E"), sub("w,t"), " \u2264 ",
        italic("E\u0302"), sub("w,t"),
        ", with the hat denoting forecast values. A small curtailment "
        "slack is admitted at the model level so that the underlying "
        "convex programme remains feasible when the renewable forecast "
        "transiently exceeds the load.",
    ], first_line_indent_cm=0.0)

    add_mixed(doc, [
        "The battery storage system is described by a discrete-time "
        "state-of-charge (SOC) dynamic equation with constant "
        "charge/discharge efficiencies. Let ",
        italic("S"), sub("t"), " \u2208 [",
        italic("S"), sub("min"), ", ", italic("S"), sub("max"), "] ",
        "denote the dimensionless state-of-charge at hour ", italic("t"),
        ", ", italic("C"), " (kWh) the rated capacity, ",
        italic("P"), sub("max"), " (kW) the rated power, ",
        italic("\u03b7"), sub("c"), " and ", italic("\u03b7"), sub("d"),
        " the charging and discharging efficiencies, and ",
        italic("\u0394t"), " (h) the time step. The SOC dynamics and "
        "operational limits are written as:",
    ], first_line_indent_cm=0.5)

    add_equation(doc, [
        italic("S"), sub("t+1"), " = ", italic("S"), sub("t"), " + (",
        italic("\u03b7"), sub("c"), "\u00a0",
        italic("P"), sub("c,t"), "\u00a0", italic("\u0394t"),
        " \u2212 ",
        italic("P"), sub("d,t"), "\u00a0", italic("\u0394t"), " / ",
        italic("\u03b7"), sub("d"), ") / ", italic("C"), ",",
    ], "2")

    add_equation(doc, [
        "0 \u2264 ", italic("P"), sub("c,t"), ", ",
        italic("P"), sub("d,t"), " \u2264 ", italic("P"), sub("max"),
        ",    ",
        italic("S"), sub("min"), " \u2264 ", italic("S"), sub("t"),
        " \u2264 ", italic("S"), sub("max"), ",",
    ], "3")

    add_mixed(doc, [
        "where ", italic("P"), sub("c,t"), " and ",
        italic("P"), sub("d,t"),
        " denote the instantaneous charge and discharge power, "
        "respectively. In our implementation the SOC operating band is "
        "set to [0.1, 0.9] to extend cycle life, ",
        italic("\u03b7"), sub("c"), " = ", italic("\u03b7"), sub("d"),
        " = 0.95, and the BES is sized at ", italic("C"),
        " = 162\u00a0MWh, ", italic("P"), sub("max"),
        " = 129.6\u00a0MW, corresponding to a C-rate of 0.8 that is "
        "representative of a state-of-the-art lithium-iron-phosphate "
        "stationary storage station.",
    ], first_line_indent_cm=0.0)

    # -----------------------------------------------------------------
    # 2.3 Deep neural forecasting model
    # -----------------------------------------------------------------
    add_h2(doc, "2.3",
        "LSTM-Based Dual-Attention Load-Forecasting Model")

    add_mixed(doc, [
        "The load forecaster, denoted ", italic("f"), sub("\u03b8"), "(\u00b7), ",
        "is a deep neural network whose parameters ",
        italic("\u03b8"), " are obtained off-line on a historical "
        "training set and subsequently loaded from the checkpoint file ",
        s("best_EModel_FeatureWeight4.pth", italic=True),
        " at the beginning of each simulation run; no online "
        "re-training is performed, so that the entire framework can be "
        "deployed on the same industrial hardware that runs the dispatch "
        "controller. Given a sliding window of ", italic("W"),
        " = 20 hours of ", italic("F"),
        "-dimensional feature vectors ",
        italic("x"), sub("t\u2212W+1"), ", \u2026, ",
        italic("x"), sub("t"),
        ", where ", italic("F"), " is automatically aligned with the "
        "runtime feature dimensionality through truncation or "
        "zero-padding of the trained weight matrices, the network "
        "outputs the next-step demand prediction ",
        italic("\u0177"), sub("t+1"), " = ",
        italic("f"), sub("\u03b8"), "(", italic("x"), sub("t\u2212W+1:t"),
        "). The architecture, illustrated in Fig.\u00a02, is composed of "
        "four interacting modules.",
    ], first_line_indent_cm=0.0)

    add_fig_caption(doc, "Fig. 2",
        "Architecture of the bidirectional LSTM forecaster "
        "EModel_FeatureWeight4: dynamic feature gating, bidirectional "
        "LSTM trunk, local temporal attention, feature attention and the "
        "two-headed reparameterised output layer.")

    add_mixed(doc, [
        s("Module 1 \u2014 Dynamic feature gating. ", bold=True),
        "The mean of the input window along the temporal axis is fed to "
        "a two-layer Gaussian-error-linear-unit (GELU) multilayer "
        "perceptron with a sigmoid output to produce per-feature gating "
        "coefficients ", italic("g"), " \u2208 [0, 1]",
        sup("F"), ", which are multiplicatively applied to the entire "
        "input window:",
    ], first_line_indent_cm=0.0)

    add_equation(doc, [
        italic("g"), " = ", italic("\u03c3"), "(", italic("W"), sub("2"),
        " \u00b7 GELU(", italic("W"), sub("1"),
        " \u00b7 mean", sub("t"), "(", italic("x"), sub("t"), ") + ",
        italic("b"), sub("1"), ") + ", italic("b"), sub("2"), "),    ",
        italic("\u0303x"), sub("t"), " = ", italic("x"), sub("t"),
        " \u2299 ", italic("g"), ",",
    ], "4")

    add_mixed(doc, [
        "where ", italic("\u03c3"),
        "(\u00b7) is the logistic sigmoid and \u2299 denotes "
        "element-wise multiplication; the rescaled tensor ",
        italic("\u0303x"), sub("t"), " \u2208 \u211d", sup("W\u00d7F"),
        " effectively suppresses noisy exogenous inputs (such as "
        "ambient humidity or wind speed) whose historical correlation "
        "with the port load is low, while preserving variables with "
        "high explanatory power (such as recent load lags and the "
        "time-of-day Fourier embedding). A Pearson-correlation-based "
        "prior is used at initialisation to bias the gate towards "
        "features whose univariate correlation with the target is above "
        "the average.",
    ], first_line_indent_cm=0.0)

    add_mixed(doc, [
        s("Module 2 \u2014 Bidirectional LSTM trunk. ", bold=True),
        "The gated window is processed by a stack of ", italic("L"),
        " = 2 bidirectional LSTM layers with hidden size ",
        italic("H"), " = 256, yielding a sequence of contextualised "
        "representations ", italic("h"), sub("1"), ", \u2026, ",
        italic("h"), sub("W"), " \u2208 \u211d", sup("2H"),
        ". Orthogonal initialisation is applied to the hidden-to-hidden "
        "weight matrices, Xavier initialisation to the input-to-hidden "
        "weight matrices, and the forget-gate bias is shifted to +1 to "
        "mitigate vanishing gradients in the early training phase. A "
        "layer-normalisation and a dropout with rate 0.1 are appended to "
        "stabilise the activations at inference time.",
    ], first_line_indent_cm=0.5)

    add_mixed(doc, [
        s("Module 3 \u2014 Dual attention. ", bold=True),
        "To compress the bidirectional hidden sequence into a "
        "fixed-length representation, two complementary attention "
        "sub-modules are combined. Local temporal attention computes "
        "scaled dot-product attention within a sliding window of length "
        "10 along the temporal axis, producing a temporal context vector ",
        italic("h"), sub("temp"),
        " by summation across time. Feature attention operates on the "
        "transposed sequence: it applies a two-layer GELU multilayer "
        "perceptron with layer normalisation, followed by a sigmoid, to "
        "the feature-by-time matrix and projects the result through a "
        "linear feature-projection layer of dimension 2",
        italic("H"), ", yielding the feature context vector ",
        italic("h"), sub("feat"),
        ". The two context vectors are concatenated to form the joint "
        "representation:",
    ], first_line_indent_cm=0.0)

    add_equation(doc, [
        italic("h"), sub("con"), " = [\u2009",
        italic("h"), sub("temp"), " ; ",
        italic("h"), sub("feat"), "\u2009] \u2208 \u211d",
        sup("4H"), ".",
    ], "5")

    add_mixed(doc, [
        s("Module 4 \u2014 Reparameterised regression head. ", bold=True),
        "The joint representation is mapped through a three-layer "
        "fully-connected head (256 \u2192 128 \u2192 2) with GELU "
        "activations, layer normalisation and dropout, producing a pair "
        "of values (", italic("\u03bc"), ", log ", italic("\u03c3"),
        sup("2"), ") = FC(", italic("h"), sub("con"),
        "). At inference time the point prediction in the normalised "
        "output space is obtained by the reparameterisation rule:",
    ], first_line_indent_cm=0.0)

    add_equation(doc, [
        italic("\u0177"), sub("scaled"), " = ", italic("\u03bc"),
        " + 0.1 \u00b7 exp(\u00bd log ", italic("\u03c3"), sup("2"),
        ") \u00b7 ", italic("\u03b5"),
        ",    ", italic("\u03b5"), " \u223c \ud4a9(0, 1),",
    ], "6")

    add_mixed(doc, [
        "after which a sklearn StandardScaler inverse-transform "
        "followed by an exp", italic("m"),
        "1(\u00b7) map (because the training objective was defined on "
        "log(1\u00a0+\u00a0", italic("y"),
        ")) recovers the demand prediction in physical units (kW). When "
        "the runtime feature dimensionality differs from the "
        "dimensionality used at training time, the helper "
        "convert_model_weights() pads or truncates the "
        "feature-importance vector, the LSTM input-to-hidden matrices "
        "and the first fully-connected layer weight matrix, so that the "
        "checkpoint can be reused without re-training. The model is "
        "invariably loaded in evaluation mode (model.eval()) and all "
        "subsequent forward passes are executed under torch.no_grad() "
        "context, which guarantees deterministic outputs at runtime.",
    ], first_line_indent_cm=0.0)

    add_mixed(doc, [
        "To amortise the cost of repeated forward passes across the five "
        "competing dispatch strategies, the forecaster is invoked exactly "
        "once for each simulated hour: a vector ",
        "{", italic("\u0177"), sub("0"), ", ",
        italic("\u0177"), sub("1"), ", \u2026, ",
        italic("\u0177"), sub("T+H\u22121"), "} of length ",
        italic("T"), " + ", italic("H"), ", where ", italic("H"),
        " is the MPC horizon, is pre-computed and cached at the beginning "
        "of the run and consumed in parallel by all subsequent dispatch "
        "experiments. On a GPU/accelerator-equipped node, mini-batch "
        "inference further lowers the kernel-launch overhead by a factor "
        "of 10\u201330, whereas on a pure CPU the LSTM is evaluated "
        "sequentially because its time-recurrent structure does not "
        "benefit from batch parallelism.",
    ], first_line_indent_cm=0.5)

    # -----------------------------------------------------------------
    # 2.4 Conformal prediction
    # -----------------------------------------------------------------
    add_h2(doc, "2.4",
        "Conformal Prediction for Calibrated Uncertainty Quantification")

    add_mixed(doc, [
        "Although the reparameterisation rule in Eq.\u00a0(6) provides a "
        "heuristic notion of predictive uncertainty, its coverage "
        "properties are not theoretically guaranteed because the "
        "underlying noise distribution is not necessarily Gaussian. To "
        "endow the forecast intervals with a finite-sample marginal "
        "coverage guarantee, a split-conformal calibration layer is "
        "wrapped around the pre-trained predictor. Let \ud49f",
        sub("cal"), " denote a calibration set drawn from the same "
        "distribution as the runtime data and disjoint from the "
        "training data; in our implementation \ud49f", sub("cal"),
        " corresponds to the [80\u202f%, 95\u202f%] portion of the "
        "historical record. For each calibration sample (",
        italic("x"), sub("i"), ", ", italic("y"), sub("i"),
        "), a non-conformity score ",
        italic("s"), sub("i"), " is computed:",
    ], first_line_indent_cm=0.0)

    add_equation(doc, [
        italic("s"), sub("i"), " = | ", italic("y"), sub("i"),
        " \u2212 ", italic("\u0177"), sub("i"), " | / max(| ",
        italic("\u0177"), sub("i"), " |, ", italic("\u03b5"), "),",
    ], "7")

    add_mixed(doc, [
        "where ", italic("\u03b5"), " = 1\u202f000\u00a0kW is a small "
        "floor that prevents the score from exploding when the point "
        "prediction is close to zero. With ", italic("m"),
        " calibration samples, the conformal quantile is obtained as "
        "the empirical \u2308(", italic("m"),
        "\u00a0+\u00a01)(1\u00a0\u2212\u00a0", italic("\u03b1"),
        ")\u2309/", italic("m"), "-th order statistic of {",
        italic("s"), sub("i"), "}:",
    ], first_line_indent_cm=0.0)

    add_equation(doc, [
        italic("\u0071\u0302"), " = Quantile",
        sub("\u2308(m+1)(1\u2212\u03b1)\u2309/m"),
        " ( {\u2009", italic("s"), sub("1"), ", ",
        italic("s"), sub("2"), ", \u2026, ",
        italic("s"), sub("m"), "\u2009} ),",
    ], "8")

    add_mixed(doc, [
        "where ", italic("\u03b1"), " \u2208 (0, 1) is the target "
        "miscoverage level; in our experiments ", italic("\u03b1"),
        " = 0.10 is adopted, corresponding to a target 90\u202f% "
        "marginal coverage. The prediction interval for any new sample "
        "with point prediction ", italic("\u0177"), " then reads:",
    ], first_line_indent_cm=0.0)

    add_equation(doc, [
        italic("I"), "(", italic("\u0177"), ") = [\u2009",
        italic("\u0177"), " \u2212 ", italic("\u0071\u0302"),
        " \u00b7 max(|", italic("\u0177"), "|, ", italic("\u03b5"),
        "), ", italic("\u0177"), " + ", italic("\u0071\u0302"),
        " \u00b7 max(|", italic("\u0177"), "|, ", italic("\u03b5"),
        ")\u2009].",
    ], "9")

    add_mixed(doc, [
        "Provided that the data are exchangeable, the resulting interval "
        "satisfies \u2119(", italic("y"), " \u2208 ", italic("I"),
        "(", italic("\u0177"), ")) \u2265 1 \u2212 ", italic("\u03b1"),
        ", which is the principal theoretical advantage of conformal "
        "calibration over heuristic Bayesian or Monte-Carlo wrappers: "
        "the guarantee holds without distributional assumptions on the "
        "residual and without any modification of the underlying "
        "deep-learning model. In our implementation, three nonconformity "
        "modes are supported: the default normalised mode of Eq.\u00a0(7), "
        "an absolute-residual mode obtained by setting ", italic("\u03b5"),
        " \u2192 \u221e in Eq.\u00a0(7), and a conformalised quantile-"
        "regression mode (CQR) in which Monte-Carlo dropout is used to "
        "obtain provisional quantile estimates that are subsequently "
        "re-calibrated by an additive correction. The robust controller "
        "of Section\u00a02.5 consumes the upper end of the interval as "
        "the worst-case demand bound, while the stochastic controller "
        "draws scenarios from the interior of the interval.",
    ], first_line_indent_cm=0.0)

    # -----------------------------------------------------------------
    # 2.5 Parametrized MPC
    # -----------------------------------------------------------------
    add_h2(doc, "2.5",
        "Parametrised Carbon-Aware Rolling-Horizon Controller")

    add_mixed(doc, [
        "At every hour ", italic("t"),
        " the controller solves a convex optimisation problem over the "
        "next ", italic("H"), " = 24 hours, applies only the first "
        "decision (", italic("P"), sub("c,t"), ", ",
        italic("P"), sub("d,t"), ", ",
        italic("P"), sub("g,t"),
        "), and rolls the horizon forward by one hour, in the spirit "
        "of classical rolling-horizon model predictive control. "
        "Crucially, all five dispatch strategies considered in this "
        "paper share a single parametrised CVXPY problem that is built "
        "once and warm-started at each step; the time-varying "
        "inputs\u2014point load forecast, renewable generation forecast, "
        "electricity price vector, dynamic carbon-intensity vector, "
        "initial SOC and running peak demand\u2014are exposed as ",
        s("cp.Parameter", italic=True),
        " objects whose ", s(".value", italic=True),
        " attribute is updated in place, while the structural problem "
        "(variables and constraint topology) remains immutable. This "
        "design eliminates the dominant cost of repeatedly compiling and "
        "disciplining convex problems and brings the per-step solve time "
        "well below the one-second wall-clock budget required to study "
        "year-long (8\u202f760-hour) scenarios.",
    ], first_line_indent_cm=0.0)

    add_mixed(doc, [
        "Decision variables of the parametrised programme are the charge "
        "and discharge power vectors ", italic("P"), sub("c"), ", ",
        italic("P"), sub("d"), " \u2208 \u211d", sub("+"),
        sup("H"), ", the grid import ",
        italic("P"), sub("g"), " \u2208 \u211d", sub("+"),
        sup("H"), " (non-negative because no reverse flow is allowed "
        "by default), the renewable utilisation ",
        italic("P"), sub("r"), " \u2208 \u211d", sub("+"),
        sup("H"), ", the SOC trajectory ", italic("S"),
        " \u2208 \u211d", sup("H+1"),
        " and the auxiliary monthly peak variable ",
        italic("P"), sub("peak"), " \u2208 \u211d", sub("+"),
        ". The feasible set is given by the SOC dynamics and "
        "operational limits of Eqs.\u00a0(2)\u2013(3), the energy balance "
        "of Eq.\u00a0(1), the renewable utilisation upper bound ",
        italic("P"), sub("r,t"), " \u2264 ", italic("\u0050\u0302"),
        sub("r,t"), " and a peak-demand epigraph constraint ",
        italic("P"), sub("g,t"), " \u2264 ", italic("P"), sub("peak"),
        ". The multi-objective cost function combines five terms:",
    ], first_line_indent_cm=0.5)

    add_equation(doc, [
        italic("J"), " = \u03a3", sub("t"), " ", italic("\u03c0"),
        sub("t"), " ", italic("P"), sub("g,t"), " + ",
        italic("\u03bb"), sub("c"), " \u00b7 \u03a3", sub("t"), " ",
        italic("\u03c6"), sub("t"), " ", italic("P"), sub("g,t"),
        " + ", italic("w"), sub("p"), " (", italic("P"), sub("peak"),
        " \u2212 ", italic("P"), sub("run"), ") \u2212 ",
        italic("\u03b1"), sub("s"), " ", italic("C"), " ",
        italic("S"), sub("H"), " + ", italic("R"), "(",
        italic("P"), sub("c"), ", ", italic("P"), sub("d"), "),",
    ], "10")

    add_mixed(doc, [
        "where ", italic("\u03c0"), sub("t"),
        " (CNY/kWh) is the time-of-use electricity tariff, ",
        italic("\u03c6"), sub("t"),
        " (tCO\u2082/MWh) is the dynamic marginal carbon intensity, ",
        italic("\u03bb"), sub("c"),
        " is the carbon-price coefficient that converts emission into "
        "monetary cost, ", italic("w"), sub("p"),
        " is the peak-charge weight, ", italic("P"), sub("run"),
        " is the running peak (the maximum grid import already realised "
        "in the current billing month), ",
        italic("\u03b1"), sub("s"),
        " is the shadow value of stored energy, and ", italic("R"),
        "(\u00b7) is a small regularisation term that suppresses "
        "spurious micro-cycling and gently encourages renewable "
        "utilisation. The peak term is formulated as a one-sided "
        "epigraph penalty so that the controller is only charged for "
        "increments above the running peak, and the storage shadow "
        "value is dynamically tied to the 24-hour mean tariff so that "
        "the optimal terminal SOC remains within a reasonable band "
        "instead of saturating at the upper limit.",
    ], first_line_indent_cm=0.0)

    add_body(doc,
        "Five dispatch strategies are obtained by selective activation "
        "of the terms of Eq.\u00a0(10).")

    add_body(doc,
        "(a) Baseline (no storage): the BES capacity is set to zero, so "
        "that all decision variables collapse and the grid import is "
        "forced to equal the residual load.",
        indent=False)

    add_body(doc,
        "(b) Rule-based peak shaving: a heuristic that charges during "
        "off-peak hours and discharges during peak hours, computed from "
        "the 24-hour mean tariff without solving any optimisation; this "
        "serves as an industry-grade reference baseline.",
        indent=False)

    add_body(doc,
        "(c) Economic MPC: full activation of the tariff, peak and "
        "regularisation terms with \u03bb_c\u00a0=\u00a00; the controller "
        "minimises the electricity bill subject to physical limits.",
        indent=False)

    add_mixed(doc, [
        "(d) Carbon-aware MPC: as (c) plus the carbon term with ",
        italic("\u03bb"), sub("c"), " = ", italic("p"), sub("c"),
        " / 1\u202f000 \u00b7 ", italic("\u03ba"), ", where ",
        italic("p"), sub("c"),
        " (CNY/tCO\u2082) is the carbon price and ", italic("\u03ba"),
        " is a user-tunable sensitivity that scales the relative "
        "weight of cost and carbon objectives.",
    ], first_line_indent_cm=0.0)

    add_body(doc,
        "(e) Robust MPC (conformal): the point forecast in the load "
        "parameter is replaced by the upper end of the conformal "
        "interval of Eq.\u00a0(9), so that the controller dispatches "
        "against a worst-case demand bound. A scalar safety factor "
        "linearly interpolates between the point forecast and the upper "
        "bound, allowing the operator to trade off the conservativeness "
        "of the dispatch against the expected cost.",
        indent=False)

    add_body(doc,
        "To guarantee robustness across the noisy and large-scale "
        "instances that arise in year-long simulations, the convex "
        "programme is solved by a three-tier solver chain. The splitting "
        "conic solver (SCS) is invoked first because it is empirically "
        "the most reliable at this problem size; if SCS fails or returns "
        "an inaccurate optimum, CLARABEL and OSQP are tried in turn. "
        "Should the entire chain fail, the controller falls back to a "
        "deterministic greedy peak-shaving heuristic that respects all "
        "physical limits but does not exploit the tariff signal "
        "explicitly, so that the simulation can continue without "
        "interruption.")

    # -----------------------------------------------------------------
    # 2.6 Carbon module
    # -----------------------------------------------------------------
    add_h2(doc, "2.6",
        "Regional Emission Factors and Dynamic Carbon Intensity")

    add_mixed(doc, [
        "Two distinct quantities related to grid emissions are used in "
        "this work and should be carefully distinguished. The average "
        "regional grid emission factor ",
        italic("\u03c6\u0304"), " (tCO\u2082/MWh) is a static, "
        "region-dependent quantity that maps the total grid import to "
        "the realised CO\u2082 emission and is taken from the latest "
        "official inventory published by the Ministry of Ecology and "
        "Environment; Table\u00a01 summarises the values used in this "
        "study for the six Chinese interconnection regions and the "
        "national average. The dynamic marginal carbon intensity ",
        italic("\u03c6"), sub("t"),
        ", on the other hand, varies hour by hour and reflects the "
        "differential composition of the marginal generation unit "
        "dispatched at each hour: during peak hours (08:00\u201312:00 "
        "and 14:00\u201318:00) the marginal unit is typically "
        "coal-fired and ", italic("\u03c6"), sub("t"),
        " is scaled upwards by a factor of 1.4 with respect to ",
        italic("\u03c6\u0304"),
        ", whereas during valley hours (00:00\u201306:00) the marginal "
        "share of nuclear, hydro and wind generation is higher and the "
        "dynamic intensity is scaled downwards by a factor of 0.6.",
    ], first_line_indent_cm=0.0)

    add_table_caption(doc, "Table 1",
        "Annual average regional grid CO\u2082 emission factors used in "
        "this study (tCO\u2082/MWh).")

    # ------------- Insert regional emission factor table -------------
    tbl = doc.add_table(rows=2, cols=7)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["North", "East", "Central", "South", "Northeast",
               "Northwest", "National Avg."]
    values  = ["0.8843", "0.7035", "0.5655", "0.3869",
               "0.6673", "0.6448", "0.5810"]
    for j, h in enumerate(headers):
        cell = tbl.cell(0, j)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(h)
        _set_run_font(run, size_pt=9, bold=True)
    for j, v in enumerate(values):
        cell = tbl.cell(1, j)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(v)
        _set_run_font(run, size_pt=9)
    # Springer-style horizontal rules: top, header bottom, table bottom
    from docx.oxml import OxmlElement as _O
    from docx.oxml.ns import qn as _qn
    tblPr = tbl._element.find(_qn('w:tblPr'))
    if tblPr is None:
        tblPr = _O('w:tblPr')
        tbl._element.insert(0, tblPr)
    borders = _O('w:tblBorders')
    for edge in ('top', 'bottom', 'insideH'):
        b = _O(f'w:{edge}')
        b.set(_qn('w:val'), 'single')
        b.set(_qn('w:sz'), '6')
        b.set(_qn('w:space'), '0')
        b.set(_qn('w:color'), '000000')
        borders.append(b)
    for edge in ('left', 'right', 'insideV'):
        b = _O(f'w:{edge}')
        b.set(_qn('w:val'), 'nil')
        borders.append(b)
    tblPr.append(borders)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)

    add_mixed(doc, [
        "The CarbonTracker module aggregates the hourly emissions into "
        "yearly totals and converts them into a monetary carbon cost "
        "using the prevailing Chinese carbon allowance price (",
        italic("p"), sub("c"), " = 120\u00a0CNY/tCO\u2082"
        " at the time of writing); for the carbon-aware MPC the "
        "resulting carbon weight is ", italic("\u03bb"), sub("c"),
        " = ", italic("p"), sub("c"), "\u00a0/\u00a01\u202f000 with an "
        "additional dimensionless sensitivity factor ",
        italic("\u03ba"),
        " (typically 1\u20133), so that the carbon-cost contribution of "
        "the objective is of the same order of magnitude as the "
        "electricity-cost contribution. The summary report further "
        "translates the total annual emission into the equivalent "
        "number of mature trees required for offset and into the "
        "equivalent number of average passenger cars displaced, in order "
        "to provide intuitive context to non-technical port "
        "stakeholders.",
    ], first_line_indent_cm=0.0)

    add_body(doc,
        "By embedding both the static regional factor and the dynamic "
        "marginal intensity in the multi-objective controller, the "
        "framework can answer two distinct questions of practical "
        "interest: (i) what is the absolute CO\u2082 footprint of the "
        "port over a representative operating year; and (ii) to what "
        "extent can the temporal redistribution of grid purchase enabled "
        "by the storage system reduce that footprint by shifting energy "
        "purchase from carbon-intensive peak hours to lower-intensity "
        "valley hours. The quantitative answers to these questions are "
        "the subject of the subsequent sections, which will be released "
        "in a forthcoming companion contribution.")

    # ---------------------------------------------------------------------
    # Save
    # ---------------------------------------------------------------------
    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    doc.save(output_path)
    print(f"[OK] Saved paper draft to: {output_path}")


# =============================================================================
if __name__ == "__main__":
    here = os.path.dirname(os.path.abspath(__file__))
    out = os.path.join(here, "Port_IES_Paper_Sections1-2.docx")
    build_paper(out)
    if os.path.exists(out):
        size = os.path.getsize(out)
        print(f"[INFO] File size: {size:,} bytes ({size/1024:.1f} KB)")
