"""
Merge Port_IES_Paper_Sections1-2.docx and Port_IES_Paper_Sections3-4.docx
into a single Port_IES_Paper_Full.docx, and embed the actual PNG figures
from Full_Model/outputs/ in front of every "Fig. N." caption.

Figure-to-PNG mapping (figures with no source PNG are concept diagrams
that live in Sections 2.1/2.3 — left as captions only):

  Fig. 1   — Overall architecture       : no PNG (concept diagram)
  Fig. 2   — LSTM forecaster diagram    : no PNG (concept diagram)
  Fig. 3   — KPI bar comparison         : 01_kpi_bars.png
  Fig. 4   — Cost decomposition         : 02_cost_breakdown.png
  Fig. 5   — KPI radar                  : 05_radar.png
  Fig. 6   — Pareto front               : 03_pareto.png, 10_pareto_front.png
  Fig. 7   — Carbon-price sensitivity   : exp04_carbon_cost_emission.png,
                                          exp04_marginal_abatement.png,
                                          exp04_cost_decomposition.png
  Fig. 8   — BESS investment economics  : exp06_bess_npv_irr.png,
                                          exp06_bess_payback.png,
                                          exp06_bess_investment_return.png,
                                          exp06_bess_total_cost.png,
                                          exp06_bess_optimal_sensitivity.png
  Fig. 9   — TOU sensitivity            : exp07_tou_absolute_cost.png,
                                          exp07_tou_savings_curve.png,
                                          exp07_tou_price_structure.png
  Fig. 10  — MPC timing                 : exp05_mpc_timing_violin.png,
                                          exp05_mpc_timing_cdf.png,
                                          exp05_mpc_timing_trace.png
  Fig. 11  — Noise robustness           : exp03_noise_absolute_cost.png,
                                          exp03_noise_degradation.png,
                                          exp03_noise_robust_advantage.png
  Fig. 12  — Extreme-day case study     : exp08_extreme_demand_error.png,
                                          exp08_extreme_error_hist.png,
                                          exp08_extreme_cumulative.png,
                                          exp08_extreme_energy_flow.png,
                                          exp08_extreme_strategy_response.png
"""
from __future__ import annotations
import copy
import os
import re

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT = os.path.dirname(os.path.abspath(__file__))
SEC12 = os.path.join(ROOT, 'Port_IES_Paper_Sections1-2.docx')
SEC34 = os.path.join(ROOT, 'Port_IES_Paper_Sections3-4.docx')
FULL  = os.path.join(ROOT, 'Port_IES_Paper_Full.docx')
OUT_DIR = os.path.join(ROOT, 'Full_Model', 'outputs')
EXP_DIR = os.path.join(OUT_DIR, 'experiments')
FIG_DIR = os.path.join(OUT_DIR, 'figures_from_manuscript')

# Page width 15.2 cm, margins 2.0 cm each side → ~11.2 cm usable.
IMG_WIDTH_CM = 11.0
IMG_WIDTH_HALF_CM = 5.4  # for side-by-side mode (not used here)


FIGURE_MAP = {
    # Fig. 1 — system architecture: borrowed from companion manuscript
    # (Figures.docx Fig. 2 "The schematic diagram of overall process").
    'Fig. 1':  [(FIG_DIR, 'Fig1_overall_process.png')],
    # Fig. 2 — LSTM forecaster diagram: stacked from Figures.docx Fig. 7
    # (LSTM Network Structure) + Fig. 8 (BiLSTM Network Structure).
    'Fig. 2':  [(FIG_DIR, 'Fig2a_LSTM_structure.png'),
                (FIG_DIR, 'Fig2b_BiLSTM_structure.png')],
    'Fig. 3':  [(OUT_DIR, '01_kpi_bars.png')],
    'Fig. 4':  [(OUT_DIR, '02_cost_breakdown.png')],
    'Fig. 5':  [(OUT_DIR, '05_radar.png')],
    'Fig. 6':  [(OUT_DIR, '03_pareto.png'),
                (OUT_DIR, '10_pareto_front.png')],
    'Fig. 7':  [(EXP_DIR, 'exp04_carbon_cost_emission.png'),
                (EXP_DIR, 'exp04_marginal_abatement.png'),
                (EXP_DIR, 'exp04_cost_decomposition.png')],
    'Fig. 8':  [(EXP_DIR, 'exp06_bess_npv_irr.png'),
                (EXP_DIR, 'exp06_bess_payback.png'),
                (EXP_DIR, 'exp06_bess_investment_return.png'),
                (EXP_DIR, 'exp06_bess_total_cost.png'),
                (EXP_DIR, 'exp06_bess_optimal_sensitivity.png')],
    'Fig. 9':  [(EXP_DIR, 'exp07_tou_absolute_cost.png'),
                (EXP_DIR, 'exp07_tou_savings_curve.png'),
                (EXP_DIR, 'exp07_tou_price_structure.png')],
    'Fig. 10': [(EXP_DIR, 'exp05_mpc_timing_violin.png'),
                (EXP_DIR, 'exp05_mpc_timing_cdf.png'),
                (EXP_DIR, 'exp05_mpc_timing_trace.png')],
    'Fig. 11': [(EXP_DIR, 'exp03_noise_absolute_cost.png'),
                (EXP_DIR, 'exp03_noise_degradation.png'),
                (EXP_DIR, 'exp03_noise_robust_advantage.png')],
    'Fig. 12': [(EXP_DIR, 'exp08_extreme_demand_error.png'),
                (EXP_DIR, 'exp08_extreme_error_hist.png'),
                (EXP_DIR, 'exp08_extreme_cumulative.png'),
                (EXP_DIR, 'exp08_extreme_energy_flow.png'),
                (EXP_DIR, 'exp08_extreme_strategy_response.png')],
}


# ---------------------------------------------------------------------------
# 1) Merge: append §3-4 body elements into §1-2 base document.
# ---------------------------------------------------------------------------
def merge_documents(base_path, append_path, out_path):
    base = Document(base_path)
    add  = Document(append_path)

    base_body = base.element.body
    # sectPr (page-properties marker) lives at the end of body; we must
    # insert new elements BEFORE it so the page layout is preserved.
    base_sectPr = base_body.find(qn('w:sectPr'))

    for child in list(add.element.body.iterchildren()):
        if child.tag == qn('w:sectPr'):
            continue  # do not duplicate the section properties
        new_child = copy.deepcopy(child)
        if base_sectPr is not None:
            base_sectPr.addprevious(new_child)
        else:
            base_body.append(new_child)

    base.save(out_path)
    print(f"[merge] {os.path.basename(base_path)} + "
          f"{os.path.basename(append_path)} -> {os.path.basename(out_path)}")


# ---------------------------------------------------------------------------
# 2) Insert images BEFORE each "Fig. N." caption paragraph.
# ---------------------------------------------------------------------------
def get_paragraph_text(p_el):
    return ''.join(t.text or '' for t in p_el.iter(qn('w:t')))


def insert_picture_paragraph(doc, after_anchor_para, image_path, width_cm):
    """Insert a new centered paragraph containing one inline image, placed
    immediately AFTER the given anchor paragraph (used for prepending all
    images of a figure cluster above its caption).

    Implementation: we add the picture at the end of the document (which
    appends a paragraph), then move that paragraph in front of the anchor.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    run = p.add_run()
    run.add_picture(image_path, width=Cm(width_cm))
    # Now move <p._element> just before the anchor paragraph.
    after_anchor_para.addprevious(p._element)
    return p


def embed_figures(doc_path):
    doc = Document(doc_path)
    body = doc.element.body

    # Walk every paragraph element in order. We collect captions to process,
    # so the iteration order is stable while we mutate.
    captions = []
    for p_el in body.iter(qn('w:p')):
        # ignore paragraphs inside tables (equation labels, KPI table cells)
        parent = p_el.getparent()
        in_table = False
        while parent is not None:
            if parent.tag == qn('w:tbl'):
                in_table = True
                break
            parent = parent.getparent()
        if in_table:
            continue

        txt = get_paragraph_text(p_el).strip()
        m = re.match(r'(Fig\.\s*\d+)\.', txt)
        if m:
            captions.append((m.group(1).replace(' ', ' '), p_el))

    inserted = 0
    for fig_label, p_el in captions:
        # Normalise to "Fig. N" form used in FIGURE_MAP keys.
        key = re.sub(r'\s+', ' ', fig_label)
        if key not in FIGURE_MAP:
            print(f"  [skip] {key} — no source PNG mapped (concept "
                  f"diagram)")
            continue
        for folder, fname in FIGURE_MAP[key]:
            img = os.path.join(folder, fname)
            if not os.path.exists(img):
                print(f"  [warn] missing image: {img}")
                continue
            insert_picture_paragraph(doc, p_el, img, IMG_WIDTH_CM)
            inserted += 1
            print(f"  [image] {key} <- {fname}")

    doc.save(doc_path)
    print(f"[embed] inserted {inserted} image(s) into "
          f"{os.path.basename(doc_path)}")


# ===========================================================================
if __name__ == '__main__':
    if not os.path.exists(SEC12):
        raise SystemExit(f"missing: {SEC12}")
    if not os.path.exists(SEC34):
        raise SystemExit(f"missing: {SEC34}")

    merge_documents(SEC12, SEC34, FULL)
    embed_figures(FULL)

    size = os.path.getsize(FULL)
    print(f"\n[OK] {FULL}")
    print(f"[INFO] file size: {size:,} bytes ({size/1024/1024:.2f} MB)")
